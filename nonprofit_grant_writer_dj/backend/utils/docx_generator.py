from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

def generate_docx(content):
    """
    Generate a DOCX file from the grant content
    
    Args:
        content (dict): A dictionary containing sections of the grant
        
    Returns:
        bytes: The DOCX file as bytes
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading(content.get('title', 'Grant Application'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add organization information
    doc.add_heading('Organization Information', level=1)
    org_info = content.get('organization_info', {})
    doc.add_paragraph(f"Name: {org_info.get('name', '')}")
    doc.add_paragraph(f"Mission: {org_info.get('mission', '')}")
    doc.add_paragraph(f"Website: {org_info.get('website', '')}")
    
    # Add executive summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(content.get('executive_summary', ''))
    
    # Add problem statement
    doc.add_heading('Problem Statement', level=1)
    doc.add_paragraph(content.get('problem_statement', ''))
    
    # Add project description
    doc.add_heading('Project Description', level=1)
    doc.add_paragraph(content.get('project_description', ''))
    
    # Add goals and objectives
    doc.add_heading('Goals and Objectives', level=1)
    # Handle string input and strip HTML tags, splitting on newlines to avoid iterating characters
    raw_goals = content.get('goals_objectives', [])
    if isinstance(raw_goals, str):
        clean_text = re.sub(r'<[^>]+>', '', raw_goals)
        goals_list = [line.strip() for line in clean_text.split('\n') if line.strip()]
    else:
        goals_list = raw_goals

    # Add each goal as a bulleted, bold run
    for goal in goals_list:
        p = doc.add_paragraph(goal, style='List Bullet')
        p.runs[0].bold = True
    
    # Add implementation plan
    doc.add_heading('Implementation Plan', level=1)
    doc.add_paragraph(content.get('implementation_plan', ''))
    
    # Add evaluation and impact
    doc.add_heading('Evaluation and Impact', level=1)
    doc.add_paragraph(content.get('evaluation', ''))
    
    # Add budget
    doc.add_heading('Budget', level=1)
    budget_items = content.get('budget', [])
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Amount'
    
    for item in budget_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('item', '')
        row_cells[1].text = item.get('description', '')
        row_cells[2].text = f"${item.get('amount', '0')}"
    
    # Add sustainability plan
    doc.add_heading('Sustainability Plan', level=1)
    doc.add_paragraph(content.get('sustainability', ''))
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(content.get('conclusion', ''))
    
    # Save document to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue() 